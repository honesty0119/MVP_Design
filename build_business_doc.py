from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = "业务设计文档.docx"
BLUE = "315EFB"
NAVY = "162033"
MUTED = "667085"
LIGHT = "F2F4F7"


def font(run, size=10.5, bold=False, color=NAVY):
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for i, (h, w) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]
        cell.width = Inches(w)
        shade(cell, LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(h), 9, True)
    for row in rows:
        cells = table.add_row().cells
        for i, (text, w) in enumerate(zip(row, widths)):
            cells[i].width = Inches(w)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            font(p.add_run(text), 8.5, False)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_h(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    font(p.add_run(text), 15 if level == 1 else 11.5, True, BLUE if level == 1 else NAVY)
    return p


def add_p(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.12
    if bold_prefix and text.startswith(bold_prefix):
        font(p.add_run(bold_prefix), 10.2, True)
        font(p.add_run(text[len(bold_prefix):]), 10.2)
    else:
        font(p.add_run(text), 10.2)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(.28)
        p.paragraph_format.first_line_indent = Inches(-.18)
        font(p.add_run(item), 9.8)


doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(.72)
sec.left_margin = sec.right_margin = Inches(.82)
sec.header_distance = sec.footer_distance = Inches(.35)

normal = doc.styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(10.2)
for name, size, before, after in (("Heading 1", 15, 12, 6), ("Heading 2", 11.5, 8, 4)):
    s = doc.styles[name]
    s.font.name = "Microsoft YaHei"
    s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(BLUE if name == "Heading 1" else NAVY)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)

header = sec.header.paragraphs[0]
font(header.add_run("CANDIDATE ASSIGNMENT · BUSINESS DESIGN"), 8.5, True, MUTED)
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(footer.add_run("市场线索到 MQL / SQL 的简化系统 · 仅使用模拟数据"), 8, False, MUTED)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(15)
p.paragraph_format.space_after = Pt(4)
font(p.add_run("市场线索到 MQL / SQL"), 25, True, NAVY)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(18)
font(p.add_run("业务设计与 MVP 边界说明"), 15, False, BLUE)
add_p(doc, "目标：用一个可运行、可解释的小闭环，验证线索从市场进入、外呼清洗、加微承接、销售跟进到 MQL/SQL 的流转。MVP 优先保证规则一致、过程留痕和可演示性，不用复杂 UI 掩盖业务缺口。")

add_h(doc, "1. 业务范围与关键假设")
add_bullets(doc, [
    "线索（Lead）是一次可运营的来源记录；客户（Customer）是完成身份归并后的主体。MVP 暂以手机号唯一近似去重，但不声称它等于真实客户唯一标识。",
    "MQL 是营销侧认可、值得持续培育的有效线索；SQL 是销售确认存在明确需求/场景并愿意推进的商机。SQL 不是“更高等级标签”，必须从 MQL 转化。",
    "负责人代表当前承接责任，不代表数据所有权。完整生产系统中，销售只看本人/团队数据，主管可看全量；Demo 不实现登录，但所有限制必须落在服务端。",
    "48 小时未跟进按“最近跟进时间；若无则按创建时间”计算，SQL 与无效线索不再进入提醒。这个口径用于行动提醒，不等同于销售 SLA 的最终定义。",
    "全部号码与记录均为模拟数据；不接入真实投放、外呼或企微账号。"
])

add_h(doc, "2. 主流程与责任边界")
add_table(doc, ["阶段", "业务动作", "进入/退出条件", "异常与责任"], [
    ("线索进入", "投放/表单/CSV 创建线索", "来源、渠道、手机号完整；手机号幂等去重", "脏数据逐行报错；不覆盖原归因"),
    ("外呼清洗", "外呼工具回传接通结果", "待外呼 → 接通/未接通；接通后判有效性", "回调签名、事件幂等、失败重试属集成层"),
    ("加微承接", "有效线索发起加微", "有效 → 待加微 → 已加微", "加微失败不等于线索无效；应单独记录原因"),
    ("私域跟进", "销售记录沟通与下一步", "已分配负责人；跟进记录只追加不覆盖", "超时提醒是运营信号，不自动改业务状态"),
    ("MQL / SQL", "按准入条件完成转化", "MQL 至少有负责人和跟进；SQL 需商机说明", "无效需原因；跨级状态被服务端拒绝"),
], [0.8, 1.55, 2.25, 1.9])

doc.add_page_break()
add_h(doc, "3. 领域模型与字段")
add_p(doc, "建模原则：把“当前状态”与“发生过什么”分开。leads 保存当前快照用于列表查询；status_history 和 follow_ups 保存不可覆盖的业务事实，避免一个 updated_at 同时代表外呼、跟进和转化。")
add_table(doc, ["对象", "核心字段", "设计理由 / 边界"], [
    ("Lead", "phone、source、channel、status、owner_id、created_at", "当前快照；手机号唯一仅是 MVP 假设"),
    ("StatusHistory", "from_status、to_status、changed_at、note", "状态变化独立留痕；生产需增加 operator_id"),
    ("FollowUp", "operator_id、content、next_action_at、created_at", "记录沟通事实及下一行动；不可覆盖历史"),
    ("User", "name、role", "演示负责人；生产由组织与权限中心提供"),
    ("CallbackEvent", "event_id、event_type、payload、received_at", "三方事件去重与问题追查；原始载荷需脱敏/限期保存"),
], [1.05, 2.7, 2.75])

add_h(doc, "4. 状态机与准入规则")
add_p(doc, "主路径：新线索 → 待外呼 → 已接通 → 有效 → 待加微 → 已加微 → MQL → SQL。未接通可以回到待外呼；任一非终态可按规则进入无效。无效重新打开属于纠错操作，Demo 强制备注，生产必须再校验主管权限。")
add_table(doc, ["规则", "服务端校验", "为什么不能只由前端控制"], [
    ("无效", "invalid_reason 必填", "其他入口/API 仍可绕过前端"),
    ("MQL", "已分配负责人 + 至少一条跟进", "防止营销标签空转、责任不清"),
    ("SQL", "当前为 MQL + 明确商机说明", "防止漏斗跨级、转化率失真"),
    ("回退/重开", "仅开放少量纠错路径并写备注", "状态回退影响业绩、SLA 与审计"),
], [1.0, 2.65, 2.85])

add_h(doc, "5. MVP 页面与操作")
add_bullets(doc, [
    "线索列表：按状态、来源、负责人筛选；同时呈现当前负责人、最近跟进和超时标记。",
    "线索详情：负责人分配、跟进追加、下一合法状态、无效/商机原因和完整状态历史。",
    "漏斗看板：总线索、有效率、MQL/SQL 转化率、超 48 小时未跟进数。",
    "新增/导入：单条新增已实现；批量导入 API 限 500 条并返回重复数与逐行错误。",
    "业务问答：支持“哪个渠道 SQL 转化率最高”“当前多少线索跟进超时”“查看各状态漏斗”等受控规则查询；返回答案时同步展示指标公式和快照口径，不执行任意 SQL。",
])

doc.add_page_break()
add_h(doc, "6. 三方系统对接设想")
add_table(doc, ["系统", "需要的接口 / 事件", "关键字段", "边界控制"], [
    ("投放平台", "线索拉取或 webhook", "external_lead_id、source、campaign、phone、occurred_at", "渠道原始 ID 幂等；归因变更需版本化"),
    ("外呼系统", "下发呼叫任务；结果回调", "event_id、lead_id、call_id、result、duration、occurred_at", "签名验真、乱序处理、重试/死信；回调不直接越级到 SQL"),
    ("SCRM / 企微", "客户匹配；加微状态回调", "external_contact_id、lead_id、add_status、sales_user_id", "加微失败与线索无效分离；员工离职需转接"),
    ("数据平台", "指标明细/维表同步", "状态事件、渠道、负责人、阶段时间", "离线分析不反写交易状态；个人字段脱敏"),
], [1.05, 1.65, 2.25, 1.55])
add_p(doc, "回调处理建议：网关验签 → 以 event_id 去重 → 原始事件落库 → 字段映射与业务校验 → 状态机执行 → 失败进入可重试队列。第三方“至少一次投递”是常态，幂等键不能只用手机号；同一手机号可能有多次合法业务事件。")

add_h(doc, "7. 指标定义与口径")
add_table(doc, ["指标", "Demo 口径", "生产口径注意事项"], [
    ("线索有效率", "有效及后续阶段数 / 总线索", "按线索进入周期做 cohort；剔除测试与重复线索"),
    ("加微率", "已加微及后续 / 有效及后续", "需约定加微失败、员工变更与跨账号归并"),
    ("MQL 转化率", "MQL+SQL / 有效及后续", "按首次达到时间统计，避免重复进入造成重复计数"),
    ("SQL 转化率", "SQL / MQL+SQL", "需冻结 SQL 定义；回退/作废应保留原事件"),
    ("平均响应时长", "MVP 暂未展示", "首次有效人工跟进 - 线索进入；区分工作时间"),
    ("超时未跟进", "非 SQL/无效且 48h 未跟进", "运营提醒口径与绩效/SLA 口径分离"),
], [1.3, 2.05, 3.15])
add_p(doc, "当前漏斗是存量快照，只适合 Demo 验证。正式经营分析应基于状态事件和 cohort 计算，否则月末新增线索尚未成熟，会拉低当期转化率；跨月转化又会污染分子。", "当前漏斗是存量快照")
add_p(doc, "自然语言查询采用意图规则映射到预定义参数化查询。这样能在 MVP 中展示 Agent / 数据查询潜力，同时避免把用户输入直接拼接成 SQL。若生产化，应增加语义层、字段白名单、数据权限、查询成本限制和结果口径解释。")

doc.add_page_break()
add_h(doc, "8. 权限、隐私与审计边界")
add_bullets(doc, [
    "销售：仅查看本人线索并追加跟进；不能修改他人记录、批量导出全量数据。",
    "主管：查看团队、分配/回收负责人、审批无效重开；关键操作必须记录操作人、前后值和原因。",
    "系统集成账号：只允许指定来源、事件类型和字段；不获得页面登录能力。",
    "手机号在页面与日志中应脱敏，导出需授权并留痕；原始回调 payload 设置保留期限，避免把审计等同于无限期保存个人数据。",
    "前端隐藏按钮不是权限控制。真实系统必须在每个查询与写接口中基于当前用户做数据范围校验。"
])

add_h(doc, "9. 异常场景与处理取舍")
add_table(doc, ["场景", "MVP", "生产化方案"], [
    ("重复线索", "手机号唯一，返回原记录", "手机号+业务线+归因窗口；客户主体与线索记录分层"),
    ("回调重复", "event_id 唯一，重复返回成功", "幂等表+业务版本；保证副作用只执行一次"),
    ("回调乱序/失败", "未完整实现", "occurred_at 与版本号判新旧；重试队列、死信与人工补偿"),
    ("并发状态变更", "SQLite 单机事务", "乐观锁/version；冲突返回最新状态，不静默覆盖"),
    ("批量导入脏数据", "单次 500 条，逐行错误", "异步任务、错误文件、配额与撤销/补偿"),
    ("口径变化", "代码内固定", "规则配置版本化；历史按当时规则回放或冻结"),
], [1.35, 2.0, 3.15])

add_h(doc, "10. MVP 取舍与迭代")
add_p(doc, "本次选择 Python 标准库 + SQLite + 原生页面，避免依赖安装阻塞演示，把有限时间用在状态机、幂等、历史留痕和指标口径上。未实现登录、复杂 RBAC、真实三方连接和异步任务，均已明确为生产化边界，而非默认“前端可信”。")
add_p(doc, "验证证据：自动化测试覆盖手机号去重、非法跨级、无效原因必填、外呼回调幂等、自然语言指标口径，以及新线索→待外呼→回调接通→有效→跟进→MQL→SQL 的完整业务链路。页面已实际验证查询答案与指标公式同时返回。")
add_table(doc, ["优先级", "下一步", "验收信号"], [
    ("P0", "登录、销售/主管权限、操作人审计、手机号脱敏", "越权测试通过；关键动作可追溯"),
    ("P0", "回调验签、重试、死信、乱序与并发控制", "重复/乱序/失败注入不造成状态错乱"),
    ("P1", "cohort 漏斗、阶段耗时、渠道成本与归因", "指标可按统一口径复算并解释"),
], [0.8, 3.5, 2.2])

doc.core_properties.title = "市场线索到 MQL/SQL：业务设计与 MVP 边界说明"
doc.core_properties.subject = "AI Coding / FDE 候选人作业"
doc.core_properties.author = "候选人"
doc.save(OUT)
print(OUT)
